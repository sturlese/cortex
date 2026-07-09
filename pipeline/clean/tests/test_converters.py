"""Deterministic converters: routing, table rendering, sheet profile, fail-fast paths.
No network and no real binaries: subprocess/httpx/genai are monkeypatched."""
import subprocess
import sys
import types

import pytest

from clean import converters
from clean.converters import _csv_rows, _render_table, _sheet_profile, extract, method_for_ext


def test_method_for_ext_routing():
    assert method_for_ext(".PDF") == "pdf"
    assert method_for_ext(".xlsx") == "sheet"
    assert method_for_ext(".docx") == "docx"
    assert method_for_ext(".pptx") == "office"
    assert method_for_ext(".ods") == "office"   # LibreOffice route: openpyxl cannot read .ods
    assert method_for_ext(".md") == "text"
    assert method_for_ext(".weird") == "text"


def test_render_table_escapes_pipes_and_newlines():
    md = _render_table([["a|b", "h2"], ["x\ny", "z"]])
    assert "a\\|b" in md
    assert "x y" in md
    assert md.splitlines()[1] == "| --- | --- |"


def test_csv_rows_skips_blank_lines(tmp_path):
    p = tmp_path / "d.csv"
    p.write_text("h1,h2\n\n1,2\n,,\n")
    (name, rows), = _csv_rows(str(p), ",")
    assert name == "Sheet1"
    assert rows == [["h1", "h2"], ["1", "2"]]


def test_sheet_profile_samples_and_counts(tmp_path):
    p = tmp_path / "big.csv"
    lines = ["col1,col2"] + [f"v{i},w{i}" for i in range(100)]
    p.write_text("\n".join(lines))
    profile = _sheet_profile(str(p))
    assert "(101 rows, 2 cols)" in profile
    assert "+75 more rows" in profile
    assert "v99" not in profile   # beyond the sample


def test_sheet_profile_xlsx(tmp_path):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "KPIs"
    ws.append(["metric", "value"])
    ws.append(["arr", 1000])
    p = tmp_path / "kpis.xlsx"
    wb.save(p)
    profile = _sheet_profile(str(p))
    assert "### KPIs (2 rows, 2 cols)" in profile
    assert "| arr | 1000 |" in profile


def test_extract_text(tmp_path):
    p = tmp_path / "note.md"
    p.write_text("hello")
    assert extract(str(p), "text") == {"method": "text", "text": "hello"}


def test_extract_pdf_uses_pdftotext(monkeypatch):
    def fake_run(cmd, capture_output, text):
        assert cmd[0] == "pdftotext"
        return types.SimpleNamespace(returncode=0, stdout="pdf text", stderr="")
    monkeypatch.setattr(subprocess, "run", fake_run)
    assert extract("/x.pdf", "pdf") == {"method": "pdf", "text": "pdf text"}


def test_pdftotext_failure_raises(monkeypatch):
    monkeypatch.setattr(subprocess, "run",
                        lambda *a, **k: types.SimpleNamespace(returncode=1, stdout="", stderr="boom"))
    with pytest.raises(RuntimeError, match="pdftotext"):
        extract("/x.pdf", "pdf")


def test_extract_office_roundtrips_through_gotenberg(monkeypatch, tmp_path):
    src = tmp_path / "deck.pptx"
    src.write_bytes(b"fake")
    monkeypatch.setattr(converters.httpx, "post",
                        lambda url, files, timeout: types.SimpleNamespace(status_code=200, content=b"%PDF"))
    monkeypatch.setattr(converters, "_pdftotext", lambda path: "converted text")
    assert extract(str(src), "office") == {"method": "office", "text": "converted text"}


def test_office_gotenberg_error_raises(monkeypatch, tmp_path):
    src = tmp_path / "deck.pptx"
    src.write_bytes(b"fake")
    monkeypatch.setattr(converters.httpx, "post",
                        lambda url, files, timeout: types.SimpleNamespace(status_code=500, text="err"))
    with pytest.raises(RuntimeError, match="gotenberg"):
        extract(str(src), "office")


def test_docx_extraction(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("Section", level=1)
    doc.add_paragraph("Body paragraph.")
    p = tmp_path / "m.docx"
    doc.save(p)
    res = extract(str(p), "docx")
    assert "## Section" in res["text"]
    assert "Body paragraph." in res["text"]


def test_vision_fails_fast_without_key(monkeypatch):
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)
    with pytest.raises(RuntimeError, match="GEMINI_API_KEY"):
        converters.vision_extract("/nonexistent.pdf")


def test_vision_extract_inline_small_pdf(monkeypatch, tmp_path):
    """Wires a fake google.genai module: small PDFs go inline, response text is returned."""
    p = tmp_path / "scan.pdf"
    p.write_bytes(b"%PDF fake")
    monkeypatch.setenv("GEMINI_API_KEY", "fake-key")

    calls = {}

    class FakeModels:
        def generate_content(self, model, contents):
            calls["model"] = model
            calls["contents"] = contents
            return types.SimpleNamespace(text="ocr result")

    class FakeClient:
        def __init__(self, api_key):
            calls["api_key"] = api_key
            self.models = FakeModels()
            self.files = types.SimpleNamespace(upload=lambda file: "uploaded")

    fake_types = types.SimpleNamespace(
        Part=types.SimpleNamespace(from_bytes=lambda data, mime_type: ("part", mime_type)))
    fake_genai = types.ModuleType("google.genai")
    fake_genai.Client = FakeClient
    fake_genai.types = fake_types
    fake_google = types.ModuleType("google")
    fake_google.genai = fake_genai
    monkeypatch.setitem(sys.modules, "google", fake_google)
    monkeypatch.setitem(sys.modules, "google.genai", fake_genai)
    monkeypatch.setitem(sys.modules, "google.genai.types", fake_types)

    res = converters.vision_extract(str(p))
    assert res["method"] == "vision"
    assert res["text"] == "ocr result"
    assert res["model"]                      # provenance present
    assert calls["contents"][0][1] == "application/pdf"   # inline part, not Files API

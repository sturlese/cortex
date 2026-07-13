"""Deterministic converters (the HANDS): pdf, sheet, docx, office, text.

Each one returns faithful text; the judgment (quality, representation, cleanup) belongs to the agent.
"""
import os
import subprocess
import tempfile

import httpx
from docx import Document
from openpyxl import load_workbook

EXT_METHOD = {
    ".pdf": "pdf",
    ".xlsx": "sheet", ".xlsm": "sheet", ".xls": "sheet", ".csv": "sheet", ".tsv": "sheet",
    ".docx": "docx",
    # .ods goes through LibreOffice (Gotenberg), NOT openpyxl — openpyxl cannot read OpenDocument
    # spreadsheets and raised on every pass, wedging the doc in an error/requeue loop forever.
    ".pptx": "office", ".ppt": "office", ".doc": "office", ".odt": "office", ".odp": "office",
    ".ods": "office", ".rtf": "office",
    ".md": "text", ".txt": "text", ".json": "text",
}

SHEET_MAX_ROWS = 5000     # hard cap on rows read per sheet
SHEET_SAMPLE_ROWS = 25    # compact profile the agent sees


def method_for_ext(ext: str) -> str:
    return EXT_METHOD.get(ext.lower(), "text")


def _pdftotext(path: str) -> str:
    r = subprocess.run(["pdftotext", "-layout", path, "-"], capture_output=True, text=True)
    if r.returncode != 0:
        raise RuntimeError(f"pdftotext rc={r.returncode}: {r.stderr[:500]}")
    return r.stdout


def _office_to_pdf(path: str) -> bytes:
    url = os.environ.get("GOTENBERG_URL", "http://gotenberg:3000")  # call-time read, never at import
    with open(path, "rb") as f:
        files = {"files": (os.path.basename(path), f.read())}
    resp = httpx.post(f"{url}/forms/libreoffice/convert", files=files, timeout=180)
    if resp.status_code != 200:
        raise RuntimeError(f"gotenberg {resp.status_code}: {resp.text[:300]}")
    return resp.content


def _docx_to_md(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "heading" in style:
            parts.append(f"## {t}")
        else:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(parts)


def _render_table(rs) -> str:
    def cell(c):
        return c.replace("|", "\\|").replace("\n", " ").strip()
    out = ["| " + " | ".join(cell(c) for c in rs[0]) + " |",
           "| " + " | ".join("---" for _ in rs[0]) + " |"]
    for r in rs[1:]:
        out.append("| " + " | ".join(cell(c) for c in r) + " |")
    return "\n".join(out)


def _xls_rows(path: str) -> list:
    """Read legacy .xls (old BIFF) with xlrd -> [(sheet_name, rows)]. openpyxl can't read .xls."""
    import xlrd
    book = xlrd.open_workbook(path)
    out = []
    for sh in book.sheets():
        rows = []
        for ri in range(min(sh.nrows, SHEET_MAX_ROWS)):
            cells = []
            for v in sh.row_values(ri):
                if v is None or v == "":
                    cells.append("")
                elif isinstance(v, float) and v.is_integer():
                    cells.append(str(int(v)))  # 5000.0 -> "5000"
                else:
                    cells.append(str(v))
            if any(c.strip() for c in cells):
                rows.append(cells)
        out.append((sh.name, rows))
    return out


def _xlsx_rows(path: str) -> list:
    wb = load_workbook(path, read_only=True, data_only=True)
    out = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if row and any(c is not None and str(c).strip() != "" for c in row):
                rows.append(["" if c is None else str(c) for c in row])
            if len(rows) >= SHEET_MAX_ROWS:
                break
        out.append((name, rows))
    return out


def _csv_rows(path: str, delim: str) -> list:
    import csv
    rows = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        for r in csv.reader(f, delimiter=delim):
            cells = ["" if c is None else str(c) for c in r]
            if any(c.strip() for c in cells):
                rows.append(cells)
            if len(rows) >= SHEET_MAX_ROWS:
                break
    return [("Sheet1", rows)]


def sheet_rows(path: str) -> list:
    """Full parsed grid per sheet: [(sheet_name, rows)] with rows capped at SHEET_MAX_ROWS.
    Shared by the page profile (below) and the facts layer (facts.py reads the same grid the
    validator re-reads — one parse, one truth)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".xls":
        return _xls_rows(path)
    if ext == ".csv":
        return _csv_rows(path, ",")
    if ext == ".tsv":
        return _csv_rows(path, "\t")
    return _xlsx_rows(path)


def _sheet_profile(path: str) -> str:
    """Compact per-sheet profile: dimensions + a sample of rows. The full grid stays in the
    original file (the page will carry `detail_in_source: true`)."""
    sheets = sheet_rows(path)
    parts = []
    for name, rows in sheets:
        if not rows:
            continue
        sample = rows[: SHEET_SAMPLE_ROWS + 1]
        parts.append(
            f"### {name} ({len(rows)} rows, {len(rows[0])} cols)\n\n{_render_table(sample)}"
            + (f"\n\n_(+{len(rows) - len(sample)} more rows in the original file)_" if len(rows) > len(sample) else "")
        )
    return "\n\n".join(parts)


VISION_OCR_PROMPT = (
    "Transcribe this document FAITHFULLY and COMPLETELY to text/Markdown. Keep tables as tables. "
    "Mark illegible passages as [illegible]. Do not summarize, invent or interpret — ONLY transcribe "
    "the content exactly as it appears."
)


def vision_extract(path: str) -> dict:
    """OCR a scanned/visual PDF via Gemini (native-PDF, single call). The engine behind the
    agent's ocr() tool. Lazy SDK import (only loaded when vision is used)."""
    model = os.environ.get("VISION_MODEL", "gemini-3-flash-preview")
    key = os.environ.get("GEMINI_API_KEY")
    if not key:
        raise RuntimeError("method=vision requires GEMINI_API_KEY (Google AI Studio); not set in the environment")
    from google import genai
    from google.genai import types
    client = genai.Client(api_key=key)
    if os.path.getsize(path) <= 14 * 1024 * 1024:
        # INLINE bytes: no Files API -> no filename header (non-ASCII filenames break the Files
        # API's ASCII header encoding). Gemini accepts native PDF inline up to ~20MB.
        with open(path, "rb") as f:
            part = types.Part.from_bytes(data=f.read(), mime_type="application/pdf")
    else:
        # large PDF -> Files API, but through a copy with an ASCII name (avoids the header issue)
        with open(path, "rb") as src, tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(src.read())
            tmpname = tmp.name
        try:
            part = client.files.upload(file=tmpname)
        finally:
            os.unlink(tmpname)
    r = client.models.generate_content(model=model, contents=[part, VISION_OCR_PROMPT])
    return {"method": "vision", "text": r.text or "", "model": model}   # model -> faithful provenance


def extract(path: str, method: str) -> dict:
    """Returns {method, text}. `text` is what the agent sees."""
    if method == "pdf":
        return {"method": "pdf", "text": _pdftotext(path)}
    if method == "office":
        pdf = _office_to_pdf(path)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(pdf)
            tmp.flush()
            return {"method": "office", "text": _pdftotext(tmp.name)}
    if method == "docx":
        return {"method": "docx", "text": _docx_to_md(path)}
    if method == "sheet":
        return {"method": "sheet", "text": _sheet_profile(path)}
    with open(path, encoding="utf-8", errors="replace") as f:
        return {"method": "text", "text": f.read()}
